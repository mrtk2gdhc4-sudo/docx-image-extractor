"""
DOCX Image Extractor + Editor + Trim Detector Service
- /health      : status check
- /extract     : extracts text + images from .docx (legacy v1)
- /edit-docx   : edit text in-place preserving images and layout
- /detect-trim : detect trim size and margins from .docx
"""

from flask import Flask, request, jsonify, send_file
from docx import Document
from docx.oxml.ns import qn
from openai import OpenAI
import zipfile
import base64
import io
import os
import re

app = Flask(__name__)

openai_client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))


# ============================================================
# HEALTH
# ============================================================
@app.route('/health', methods=['GET'])
def health():
    has_openai = bool(os.environ.get('OPENAI_API_KEY'))
    return jsonify({
        'status': 'ok',
        'service': 'docx-image-extractor',
        'openai_configured': has_openai,
        'endpoints': ['/health', '/extract', '/edit-docx', '/detect-trim']
    })


# ============================================================
# LEGACY EXTRACT (for v1 illustrated workflow)
# ============================================================
@app.route('/extract', methods=['POST'])
def extract_docx():
    try:
        docx_bytes = None
        filename = 'unknown.docx'

        if 'file' in request.files:
            uploaded_file = request.files['file']
            docx_bytes = uploaded_file.read()
            filename = uploaded_file.filename or 'upload.docx'
        elif request.is_json:
            data = request.json
            if not data or 'file_base64' not in data:
                return jsonify({'error': 'Missing file_base64 or file upload'}), 400
            docx_bytes = base64.b64decode(data['file_base64'])
            filename = data.get('filename', 'unknown.docx')
        else:
            return jsonify({'error': 'No file provided'}), 400

        if not docx_bytes or len(docx_bytes) < 100:
            return jsonify({'error': f'File too small'}), 400

        docx_io = io.BytesIO(docx_bytes)
        images = []
        image_filenames_in_order = []

        try:
            with zipfile.ZipFile(docx_io, 'r') as z:
                media_files = sorted([f for f in z.namelist() if f.startswith('word/media/')])
                for media_path in media_files:
                    image_data = z.read(media_path)
                    image_name = os.path.basename(media_path)
                    images.append({
                        'name': image_name,
                        'original_path': media_path,
                        'data_base64': base64.b64encode(image_data).decode('utf-8'),
                        'mime_type': guess_mime(image_name),
                        'size_bytes': len(image_data)
                    })
                    image_filenames_in_order.append(image_name)
        except zipfile.BadZipFile:
            return jsonify({'error': 'File is not a valid .docx'}), 400

        docx_io.seek(0)
        doc = Document(docx_io)
        text_parts = []
        image_counter = 0
        marker_to_image = {}

        for para in doc.paragraphs:
            para_has_image = False
            for run in para.runs:
                drawings = run.element.findall('.//' + qn('w:drawing'))
                for _ in drawings:
                    image_counter += 1
                    marker = f'[IMAGE_{image_counter:03d}]'
                    text_parts.append(marker)
                    if image_counter <= len(image_filenames_in_order):
                        marker_to_image[marker] = image_filenames_in_order[image_counter - 1]
                    para_has_image = True
            if para.text.strip():
                text_parts.append(para.text)
            elif not para_has_image:
                text_parts.append('')

        full_text = '\n\n'.join(text_parts)
        full_text = re.sub(r'\n{3,}', '\n\n', full_text).strip()

        return jsonify({
            'text_with_markers': full_text,
            'images': images,
            'image_count': image_counter,
            'marker_to_image': marker_to_image,
            'filename': filename,
            'received_bytes': len(docx_bytes),
            'success': True
        })

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc(), 'success': False}), 500


# ============================================================
# EDIT IN PLACE (for v2 illustrated workflow — preserves layout)
# ============================================================
EDIT_SYSTEM_PROMPT = """You are a professional copyeditor. You will receive paragraphs of book text, one paragraph per line, numbered. For EACH numbered paragraph, return the edited version with the same number.

EDITING RULES:
- Fix grammar, punctuation, capitalization, hyphenation, obvious typos
- Preserve the author's voice exactly (sentence fragments, polysyndeton, repetition for emphasis)
- Preserve dialogue exactly (only fix capitalization at start of quotes)
- Preserve all paragraph numbers — never merge or split paragraphs
- If a paragraph is already correct, return it unchanged
- If a paragraph is empty or just whitespace, return empty
- DO NOT add commentary, headers, or markdown
- DO NOT rewrite for style

OUTPUT FORMAT (critical):
Return paragraphs in this exact format:
[1] edited paragraph 1 text
[2] edited paragraph 2 text
[3] edited paragraph 3 text

One paragraph per line. The number in [N] must match the input number exactly.
"""


def edit_paragraphs_batch(paragraphs_with_indices):
    if not paragraphs_with_indices:
        return {}

    lines = []
    for idx, text in paragraphs_with_indices:
        clean_text = text.replace('\n', ' ').replace('\r', ' ').strip()
        lines.append(f'[{idx}] {clean_text}')
    user_message = '\n'.join(lines)

    response = openai_client.chat.completions.create(
        model='gpt-4o',
        temperature=0.3,
        max_tokens=4000,
        messages=[
            {'role': 'system', 'content': EDIT_SYSTEM_PROMPT},
            {'role': 'user', 'content': user_message},
        ],
    )
    edited_text = response.choices[0].message.content.strip()

    result = {}
    for line in edited_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        match = re.match(r'^\[(\d+)\]\s*(.*)$', line)
        if match:
            result[int(match.group(1))] = match.group(2)

    return result


def update_paragraph_text(para, new_text):
    text_only_runs = []
    for run in para.runs:
        has_drawing = run.element.findall('.//' + qn('w:drawing'))
        if not has_drawing:
            text_only_runs.append(run)

    if not text_only_runs:
        return

    for run in text_only_runs:
        run.text = ''
    text_only_runs[0].text = new_text


@app.route('/edit-docx', methods=['POST'])
def edit_docx():
    try:
        if not os.environ.get('OPENAI_API_KEY'):
            return jsonify({'error': 'OPENAI_API_KEY not configured on server'}), 500

        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        uploaded_file = request.files['file']
        docx_bytes = uploaded_file.read()
        filename = uploaded_file.filename or 'upload.docx'

        if not docx_bytes or len(docx_bytes) < 100:
            return jsonify({'error': 'File too small'}), 400

        docx_io = io.BytesIO(docx_bytes)
        doc = Document(docx_io)

        paragraphs_to_edit = []
        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text or len(text) < 3:
                continue
            paragraphs_to_edit.append((idx, text))

        total_paragraphs = len(paragraphs_to_edit)
        if total_paragraphs == 0:
            return jsonify({'error': 'No editable paragraphs found'}), 400

        BATCH_SIZE = 30
        all_edits = {}
        batch_count = 0

        for batch_start in range(0, total_paragraphs, BATCH_SIZE):
            batch = paragraphs_to_edit[batch_start:batch_start + BATCH_SIZE]
            batch_count += 1
            try:
                edits = edit_paragraphs_batch(batch)
                all_edits.update(edits)
            except Exception as e:
                return jsonify({
                    'error': f'GPT-4o failed on batch {batch_count}: {str(e)}',
                    'batches_completed': batch_count - 1,
                    'paragraphs_attempted': batch_start + len(batch)
                }), 500

        edits_applied = 0
        edits_skipped = 0
        for idx, original_text in paragraphs_to_edit:
            if idx in all_edits:
                edited_text = all_edits[idx]
                ratio = len(edited_text) / max(len(original_text), 1)
                if 0.5 <= ratio <= 1.8:
                    update_paragraph_text(doc.paragraphs[idx], edited_text)
                    edits_applied += 1
                else:
                    edits_skipped += 1
            else:
                edits_skipped += 1

        output_io = io.BytesIO()
        doc.save(output_io)
        output_io.seek(0)
        output_bytes = output_io.getvalue()

        return send_file(
            io.BytesIO(output_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'edited_{filename}',
        ), 200, {
            'X-Edits-Applied': str(edits_applied),
            'X-Edits-Skipped': str(edits_skipped),
            'X-Total-Paragraphs': str(total_paragraphs),
            'X-Batches': str(batch_count),
        }

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'traceback': traceback.format_exc()}), 500


# ============================================================
# TRIM SIZE DETECTION
# ============================================================
@app.route('/detect-trim', methods=['POST'])
def detect_trim():
    """
    POST /detect-trim
    Accepts a .docx via multipart upload (field 'file').
    Returns trim size and margins in inches.
    """
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        uploaded_file = request.files['file']
        docx_bytes = uploaded_file.read()
        filename = uploaded_file.filename or 'upload.docx'

        if not docx_bytes or len(docx_bytes) < 100:
            return jsonify({'error': 'File too small'}), 400

        docx_io = io.BytesIO(docx_bytes)
        doc = Document(docx_io)
        section = doc.sections[0]

        def emu_to_inches(emu_value):
            if emu_value is None:
                return None
            return round(emu_value / 914400, 3)

        trim_width = emu_to_inches(section.page_width)
        trim_height = emu_to_inches(section.page_height)
        margin_top = emu_to_inches(section.top_margin)
        margin_bottom = emu_to_inches(section.bottom_margin)
        margin_left = emu_to_inches(section.left_margin)
        margin_right = emu_to_inches(section.right_margin)
        gutter = emu_to_inches(section.gutter) if hasattr(section, 'gutter') else 0

        margin_inside = round((margin_left or 0) + (gutter or 0), 3)
        margin_outside = margin_right or 0

        common_sizes = [
            (4.25, 6.87),
            (5.0, 8.0),
            (5.06, 7.81),
            (5.25, 8.0),
            (5.5, 8.5),
            (6.0, 9.0),
            (6.14, 9.21),
            (7.0, 10.0),
            (8.5, 11.0),
        ]

        snapped = False
        original_w, original_h = trim_width, trim_height
        for w, h in common_sizes:
            if abs(trim_width - w) < 0.05 and abs(trim_height - h) < 0.05:
                trim_width = w
                trim_height = h
                snapped = True
                break

        return jsonify({
            'success': True,
            'filename': filename,
            'trim_width': trim_width,
            'trim_height': trim_height,
            'margin_top': margin_top,
            'margin_bottom': margin_bottom,
            'margin_left': margin_left,
            'margin_right': margin_right,
            'margin_inside': margin_inside,
            'margin_outside': margin_outside,
            'gutter': gutter,
            'snapped_to_standard': snapped,
            'detected_size_label': f'{trim_width} x {trim_height} inches',
            'original_measurements': {
                'width': original_w,
                'height': original_h
            }
        })

    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


# ============================================================
# HELPERS
# ============================================================
def guess_mime(filename):
    ext = filename.lower().rsplit('.', 1)[-1]
    return {
        'png': 'image/png', 'jpg': 'image/jpeg', 'jpeg': 'image/jpeg',
        'gif': 'image/gif', 'webp': 'image/webp', 'bmp': 'image/bmp',
        'tiff': 'image/tiff', 'tif': 'image/tiff',
    }.get(ext, 'application/octet-stream')


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 8080))
    app.run(host='0.0.0.0', port=port)
